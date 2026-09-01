"""
Confluence page ingester.

Uses the Confluence REST API (atlassian-python-api): list pages per space with pagination,
expand body.storage, convert HTML to plain text with BeautifulSoup, optional title prefix for RAG.
"""

from __future__ import annotations

import mimetypes
import time
from io import BytesIO
from typing import Any
from urllib.parse import urlparse

import requests

from knowledge_transfer_agent.config import get_settings
from knowledge_transfer_agent.ingestion.base import BaseIngester, Document, IngestionResult
from knowledge_transfer_agent.logging_config import get_logger

logger = get_logger(__name__)

_DEFAULT_ATTACHMENT_MAX_BYTES = 15 * 1024 * 1024  # 15MB
_DEFAULT_ATTACHMENT_LIMIT_PER_PAGE = 50
_SUPPORTED_ATTACHMENT_EXTS = {".pdf", ".docx", ".txt", ".md", ".rst"}


def _normalize_confluence_url(url: str) -> str:
    """Strip whitespace and trailing slashes; keep user-provided path (often /wiki)."""
    u = url.strip().rstrip("/")
    if not u:
        return u
    parsed = urlparse(u)
    if not parsed.scheme or not parsed.netloc:
        logger.warning("CONFLUENCE_URL may be invalid: %s", url[:80])
    return u


def _storage_html_to_text(html: str) -> str:
    """
    Confluence body.storage is HTML. Parse with BeautifulSoup for clean plain text;
    fall back to regex if beautifulsoup4 is not installed.
    """
    if not html or not html.strip():
        return ""
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text(separator=" ")
    except ImportError:
        import re

        text = re.sub(r"<[^>]+>", " ", html)
        logger.warning(
            "beautifulsoup4 not installed; using regex HTML strip. "
            "Install with: pip install beautifulsoup4"
        )
    return " ".join(text.split())


def _unwrap_results(raw: Any) -> list[dict[str, Any]]:
    """Handle list responses or {'results': [...]} from REST wrappers."""
    if raw is None:
        return []
    if isinstance(raw, list):
        return [x for x in raw if isinstance(x, dict)]
    if isinstance(raw, dict):
        results = raw.get("results")
        if isinstance(results, list):
            return [x for x in results if isinstance(x, dict)]
    return []


def _fetch_all_pages_in_space(
    confluence: Any,
    space_key: str,
    *,
    batch_size: int,
    request_delay_seconds: float,
) -> list[dict[str, Any]]:
    """Paginate get_all_pages_from_space until no more pages."""
    all_pages: list[dict[str, Any]] = []
    start = 0
    while True:
        try:
            raw = confluence.get_all_pages_from_space(
                space_key,
                expand="body.storage,version,space,history",
                start=start,
                limit=batch_size,
            )
        except Exception as e:
            logger.exception("Confluence list pages failed at start=%s space=%s", start, space_key)
            raise

        batch = _unwrap_results(raw)
        if not batch:
            break
        all_pages.extend(batch)
        if len(batch) < batch_size:
            break
        start += batch_size
        if request_delay_seconds > 0:
            time.sleep(request_delay_seconds)

    return all_pages


def _build_confluence_client(settings: Any) -> Any:
    """Construct atlassian Confluence client (Cloud token or username+token)."""
    from atlassian import Confluence

    url = _normalize_confluence_url(settings.confluence_url or "")
    token = settings.confluence_token
    cloud = settings.confluence_cloud
    username = settings.confluence_username

    if username and token:
        return Confluence(url=url, username=username, password=token, cloud=cloud)
    if token:
        # Many library versions support PAT-only for Cloud
        try:
            return Confluence(url=url, token=token, cloud=cloud)
        except TypeError:
            return Confluence(url=url, username="", password=token, cloud=cloud)
    raise ValueError("CONFLUENCE_TOKEN is required")


def _page_to_plain_text(page: dict[str, Any]) -> str:
    content = page.get("body", {}).get("storage", {}).get("value", "")
    return _storage_html_to_text(content)


def _document_text_for_index(title: str, body_text: str) -> str:
    """Prefix title so embeddings and chunks retain page context."""
    title = (title or "").strip()
    if title and body_text:
        return f"{title}\n\n{body_text}"
    return body_text or title


def _attachment_download_url(base_url: str, attachment: dict[str, Any]) -> str | None:
    """
    Build absolute download URL from Confluence attachment object.
    Atlassian responses usually include: attachment['_links']['download'] (relative).
    """
    links = attachment.get("_links") if isinstance(attachment.get("_links"), dict) else {}
    rel = links.get("download")
    if not isinstance(rel, str) or not rel.strip():
        return None
    if rel.startswith("http://") or rel.startswith("https://"):
        return rel
    return f"{base_url}{rel}"


def _download_bytes(
    url: str,
    *,
    username: str | None,
    token: str,
    timeout_seconds: int = 60,
    max_bytes: int = _DEFAULT_ATTACHMENT_MAX_BYTES,
) -> bytes:
    """
    Download binary content with Confluence auth.

    - If username is set: use Basic Auth (username, token) (typical Atlassian Cloud API token flow).
    - Else: try Bearer token.
    """
    auth = (username, token) if username else None
    headers = {"User-Agent": "knowledge-transfer-agent/1.0"}
    if not username:
        headers["Authorization"] = f"Bearer {token}"

    with requests.get(url, headers=headers, auth=auth, stream=True, timeout=timeout_seconds) as r:
        r.raise_for_status()
        chunks: list[bytes] = []
        total = 0
        for part in r.iter_content(chunk_size=1024 * 256):
            if not part:
                continue
            total += len(part)
            if total > max_bytes:
                raise ValueError(f"Attachment too large: {total} bytes > {max_bytes} bytes")
            chunks.append(part)
        return b"".join(chunks)


def _extract_text_from_attachment(
    content: bytes,
    *,
    filename: str,
) -> list[tuple[str, dict[str, Any]]]:
    """
    Extract text from an attachment.

    Returns a list of (text, metadata) items. For PDFs, returns one item per page.
    For other supported files, returns a single item.
    """
    ext = (filename or "").lower()
    if "." in ext:
        ext = "." + ext.split(".")[-1]

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as e:
            raise ImportError("pypdf required for PDF attachments") from e

        reader = PdfReader(BytesIO(content))
        out: list[tuple[str, dict[str, Any]]] = []
        page_count = len(reader.pages)
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            out.append(
                (
                    text,
                    {
                        "page_number": i + 1,
                        "page_count": page_count,
                        "file_type": ".pdf",
                    },
                )
            )
        return out

    if ext == ".docx":
        try:
            import docx  # python-docx
        except ImportError as e:
            raise ImportError("python-docx required for DOCX attachments") from e
        d = docx.Document(BytesIO(content))
        text = "\n".join(p.text for p in d.paragraphs if (p.text or "").strip()).strip()
        return [(text, {"file_type": ".docx"})] if text else []

    if ext in {".txt", ".md", ".rst"}:
        try:
            text = content.decode("utf-8", errors="replace").strip()
        except Exception:
            text = ""
        return [(text, {"file_type": ext})] if text else []

    return []


class ConfluenceIngester(BaseIngester):
    """Ingests Confluence wiki pages from configured space keys (REST API, full pagination)."""

    source_type = "confluence"

    def __init__(
        self,
        url: str | None = None,
        token: str | None = None,
        space_keys: list[str] | None = None,
    ) -> None:
        settings = get_settings()
        self.url = url or settings.confluence_url
        self.token = token or settings.confluence_token
        if space_keys is not None:
            self.space_keys = [s.strip() for s in space_keys if s.strip()]
        elif settings.confluence_space_keys:
            self.space_keys = [
                s.strip() for s in settings.confluence_space_keys.split(",") if s.strip()
            ]
        else:
            self.space_keys = []

    def ingest(self) -> IngestionResult:
        """
        Ingest Confluence pages from configured spaces.

        Requires CONFLUENCE_URL, CONFLUENCE_TOKEN, and CONFLUENCE_SPACE_KEYS.
        For Atlassian Cloud, set CONFLUENCE_USERNAME to your Atlassian email if token-only auth fails.
        Install: pip install atlassian-python-api
        """
        if not all([self.url, self.token, self.space_keys]):
            logger.warning(
                "Confluence not configured. Set CONFLUENCE_URL, CONFLUENCE_TOKEN, "
                "and CONFLUENCE_SPACE_KEYS."
            )
            return IngestionResult(
                success=False,
                documents_processed=0,
                documents_failed=0,
                source="confluence",
                source_type=self.source_type,
                documents=[],
                errors=["Confluence credentials or space keys not configured"],
            )

        try:
            from atlassian import Confluence  # noqa: F401
        except ImportError:
            return IngestionResult(
                success=False,
                documents_processed=0,
                documents_failed=0,
                source="confluence",
                source_type=self.source_type,
                documents=[],
                errors=["atlassian-python-api required: pip install atlassian-python-api"],
            )

        settings = get_settings()
        documents: list[Document] = []
        errors: list[str] = []

        try:
            confluence = _build_confluence_client(settings)
        except Exception as e:
            return IngestionResult(
                success=False,
                documents_processed=0,
                documents_failed=1,
                source="confluence",
                source_type=self.source_type,
                documents=[],
                errors=[f"Confluence client failed: {e}"],
            )

        batch_size = max(1, min(settings.confluence_page_batch_size, 250))
        delay = float(settings.confluence_request_delay_seconds or 0.0)

        base_url = _normalize_confluence_url(self.url or "")
        username = settings.confluence_username
        token = settings.confluence_token or ""

        try:
            for space_key in self.space_keys:
                try:
                    pages = _fetch_all_pages_in_space(
                        confluence,
                        space_key,
                        batch_size=batch_size,
                        request_delay_seconds=float(delay),
                    )
                    logger.info(
                        "Confluence space %s: fetched %d pages",
                        space_key,
                        len(pages),
                    )

                    for page in pages:
                        try:
                            title = page.get("title") or ""
                            body_text = _page_to_plain_text(page)
                            if not body_text.strip() and not title.strip():
                                continue

                            indexed_text = _document_text_for_index(title, body_text)
                            page_id = page.get("id")
                            version = page.get("version") or {}
                            hist = page.get("history") or {}
                            last_upd = None
                            if isinstance(hist.get("lastUpdated"), dict):
                                last_upd = hist.get("lastUpdated", {}).get("when")

                            doc = Document(
                                content=indexed_text,
                                metadata={
                                    "page_id": page_id,
                                    "title": title,
                                    "space_key": space_key,
                                    "version": version.get("number") if isinstance(version, dict) else version,
                                    "last_updated": last_upd,
                                    "content_kind": "confluence_page",
                                },
                                source=f"{base_url}/pages/viewpage.action?pageId={page_id}",
                                source_type=self.source_type,
                                doc_id=str(page_id) if page_id is not None else None,
                            )
                            documents.append(doc)

                            # --- Attachments (production: uploaded docs referenced from Confluence pages) ---
                            if page_id is not None:
                                try:
                                    raw_atts = confluence.get_attachments_from_content(
                                        page_id,
                                        start=0,
                                        limit=_DEFAULT_ATTACHMENT_LIMIT_PER_PAGE,
                                        expand="version,metadata,extensions",
                                    )
                                    attachments = _unwrap_results(raw_atts)
                                except Exception as e:
                                    errors.append(f"Page {page_id} attachments: {e!s}")
                                    attachments = []

                                for att in attachments:
                                    try:
                                        att_id = att.get("id")
                                        att_title = (att.get("title") or "").strip()
                                        if not att_title:
                                            continue
                                        ext = (("." + att_title.split(".")[-1]).lower()) if "." in att_title else ""
                                        if ext and ext not in _SUPPORTED_ATTACHMENT_EXTS:
                                            continue

                                        dl_url = _attachment_download_url(base_url, att)
                                        if not dl_url:
                                            continue

                                        data = _download_bytes(
                                            dl_url,
                                            username=username,
                                            token=token,
                                            timeout_seconds=60,
                                            max_bytes=_DEFAULT_ATTACHMENT_MAX_BYTES,
                                        )
                                        extracted = _extract_text_from_attachment(data, filename=att_title)
                                        if not extracted:
                                            continue

                                        # One Document per extracted unit (PDF pages become multiple docs)
                                        for unit_text, unit_meta in extracted:
                                            unit_text = (unit_text or "").strip()
                                            if not unit_text:
                                                continue
                                            base_meta = {
                                                "content_kind": "confluence_attachment",
                                                "page_id": page_id,
                                                "page_title": title,
                                                "space_key": space_key,
                                                "attachment_id": att_id,
                                                "attachment_name": att_title,
                                                "attachment_download_url": dl_url,
                                            }
                                            base_meta.update(unit_meta or {})
                                            # Stable id for citations; include page_number when present
                                            pn = base_meta.get("page_number")
                                            suffix = f"#p{pn}" if isinstance(pn, int) else ""
                                            doc_id = (
                                                f"{page_id}:{att_id}:{att_title}{suffix}"
                                                if att_id is not None
                                                else f"{page_id}:{att_title}{suffix}"
                                            )
                                            documents.append(
                                                Document(
                                                    content=f"{att_title}\n\n{unit_text}",
                                                    metadata=base_meta,
                                                    source=dl_url,
                                                    source_type=self.source_type,
                                                    doc_id=doc_id,
                                                )
                                            )

                                    except ImportError as e:
                                        errors.append(f"Page {page_id} attachment {att.get('title')}: {e!s}")
                                    except Exception as e:
                                        errors.append(f"Page {page_id} attachment {att.get('title')}: {e!s}")

                        except Exception as e:
                            errors.append(f"Page {page.get('id')}: {e!s}")

                except Exception as e:
                    errors.append(f"Space {space_key}: {e!s}")
                    logger.exception("Confluence space ingest failed: %s", space_key)

            return IngestionResult(
                success=len(documents) > 0,
                documents_processed=len(documents),
                documents_failed=len(errors),
                source="confluence",
                source_type=self.source_type,
                documents=documents,
                errors=errors if errors else [],
                metadata={
                    "spaces": self.space_keys,
                    "page_count": len(documents),
                },
            )

        except Exception as e:
            logger.exception("Confluence ingestion failed")
            return IngestionResult(
                success=False,
                documents_processed=len(documents),
                documents_failed=1,
                source="confluence",
                source_type=self.source_type,
                documents=documents,
                errors=[str(e)],
            )
